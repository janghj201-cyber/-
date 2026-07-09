"""
위베이프 일일 보고 스크립트 v4.0
경로: staff_todos/{이름}_{날짜} / handover/{매장}_{날짜} / checks/{매장}/{날짜}/clean
v4 추가:
  - 청소 체크 타임스탬프 기반 "몰아치기" 패턴 감지
  - 인수인계 수신확인(다른 사람 인계를 내가 확인한 건수) 집계
  - Claude API를 통한 업무/인수인계 내용 질적 판단
"""
import os, json, urllib.request, urllib.error, smtplib, io
from datetime import datetime, timedelta
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.base import MIMEBase
from email import encoders

FIREBASE_API_KEY   = os.environ["FIREBASE_API_KEY"]
GMAIL_USER         = os.environ["GMAIL_USER"]
GMAIL_APP_PASSWORD = os.environ["GMAIL_APP_PASSWORD"]
REPORT_TO          = os.environ["REPORT_TO_EMAIL"]
ANTHROPIC_API_KEY  = os.environ.get("ANTHROPIC_API_KEY")  # 없으면 LLM 판단 단계는 건너뜀
PROJECT_ID         = "wevape-schedule"
BASE_URL           = f"https://firestore.googleapis.com/v1/projects/{PROJECT_ID}/databases/(default)/documents"

STORES = [
    {"id":"yeonsu",    "name":"인천 연수점",   "short":"연수점"},
    {"id":"nonhyeon",  "name":"인천 논현점",   "short":"논현점"},
    {"id":"rodeo",     "name":"구월 로데오점", "short":"로데오점"},
    {"id":"gilbyeong", "name":"구월 길병원점", "short":"길병원점"},
    {"id":"airport",   "name":"인천공항점",    "short":"공항점"},
    {"id":"geomdan",   "name":"검단점",        "short":"검단점"},
    {"id":"gyesan",    "name":"계산점",        "short":"계산점"},
    {"id":"sangdong",  "name":"부천 상동점",   "short":"상동점"},
    {"id":"sijungdong","name":"부천 신중동점", "short":"신중동점"},
]

STAFF = ["오명록","고아현","장현진","장대운","신재현","정희경",
         "조효정","홍다운","이종혁","원주현","김형진","윤하람",
         "차영근","정유진","안태민","김다정"]

def fb_get(path):
    try:
        req = urllib.request.Request(f"{BASE_URL}/{path}?key={FIREBASE_API_KEY}")
        with urllib.request.urlopen(req, timeout=15) as r:
            d = json.load(r)
            return parse_fields(d["fields"]) if "fields" in d else None
    except:
        return None

def parse_value(v):
    if "stringValue"  in v: return v["stringValue"]
    if "booleanValue" in v: return v["booleanValue"]
    if "integerValue" in v: return int(v["integerValue"])
    if "doubleValue"  in v: return v["doubleValue"]
    if "arrayValue"   in v: return [parse_value(x) for x in v["arrayValue"].get("values",[])]
    if "mapValue"     in v: return parse_fields(v["mapValue"].get("fields",{}))
    return None

def parse_fields(f): return {k: parse_value(v) for k,v in f.items()}

def get_todos(name, dk):
    d = fb_get(f"staff_todos/{name}_{dk}"); return (d.get("items") or []) if d else []

# 청소 체크값 판독 (구버전 boolean / 신버전 {v,t} 객체 모두 지원)
def ck_val(x):
    if isinstance(x, dict): return x.get("v") is True
    return x is True
def ck_time(x):
    if isinstance(x, dict): return x.get("t")
    return None

def get_clean(sid, dk):
    d = fb_get(f"checks/{sid}/{dk}/clean")
    if not d: return 0, 0, []
    keys = [k for k in d if k.startswith("item_")]
    total = len(keys) or 6
    done = sum(1 for k in keys if ck_val(d.get(k)))
    times = sorted([ck_time(d.get(k)) for k in keys if ck_time(d.get(k))])
    return done, total, times

def check_clean_pattern(times, total):
    """청소 체크 타임스탬프가 비정상적으로 몰려서 찍혔는지(일괄 클릭 의심) 판단"""
    if len(times) < max(3, total // 2):
        return None
    span_sec = (max(times) - min(times)) / 1000
    if len(times) >= 4 and span_sec < 60:
        return f"청소 체크 {len(times)}건이 {int(span_sec)}초 안에 몰려서 처리됨 (일괄 클릭 의심)"
    return None

def get_handover(sid, dk):
    d = fb_get(f"handover/{sid}_{dk}"); return (d.get("items") or []) if d else []

def get_handover_confirmations(dk):
    """해당 날짜, 전 매장 기준 - 각 직원이 '확인 처리'한 인수인계 건수 (수신 성실도)"""
    counts = {}
    for s in STORES:
        for it in get_handover(s["id"], dk):
            if it.get("confirmed") and it.get("confirmedBy"):
                who = it["confirmedBy"]
                counts[who] = counts.get(who, 0) + 1
    return counts

def get_dayoff():
    d = fb_get("config/dayoff"); return d or {}
def get_projects(name):
    d = fb_get(f"staff_projects/{name}"); return (d.get("items") or []) if d else []

# ── Claude API를 통한 업무/인수인계 내용 질적 판단 ──────────────
def llm_judge(staff_texts):
    """
    staff_texts: {이름: {"todos": [...], "handovers": [...]}}
    반환: (judgments dict, error string or None)
    """
    if not ANTHROPIC_API_KEY:
        return {}, None
    lines = []
    for name, d in staff_texts.items():
        if not d["todos"] and not d["handovers"]:
            continue
        lines.append(f"[{name}]")
        if d["todos"]:
            lines.append("업무: " + " / ".join(d["todos"][:8]))
        if d["handovers"]:
            lines.append("인수인계: " + " / ".join(d["handovers"][:5]))
    if not lines:
        return {}, None

    prompt = (
        "다음은 위베이프 매장 직원들이 하루 동안 입력한 '업무 내용'과 '인수인계 내용'입니다.\n"
        "각 직원별로 입력 내용이 구체적이고 실질적인지, 아니면 형식적이거나 성의없이 짧게만 썼는지를 "
        "한 줄로 짧고 담백하게 평가해 주세요. 과장하지 말고 사실 기반으로만 판단하세요.\n"
        "반드시 아래 JSON 형식으로만 응답하세요 (다른 설명, 코드블록 없이 순수 JSON만):\n"
        '{"직원이름": "한 줄 평가"}\n\n'
        + "\n".join(lines)
    )
    try:
        body = json.dumps({
            "model": "claude-sonnet-5",
            "max_tokens": 1000,
            "messages": [{"role": "user", "content": prompt}]
        }).encode()
        req = urllib.request.Request(
            "https://api.anthropic.com/v1/messages",
            data=body, method="POST"
        )
        req.add_header("x-api-key", ANTHROPIC_API_KEY)
        req.add_header("anthropic-version", "2023-06-01")
        req.add_header("content-type", "application/json")
        with urllib.request.urlopen(req, timeout=30) as r:
            resp = json.load(r)
        text_block = next((b.get("text") for b in resp.get("content", []) if b.get("type") == "text"), None)
        if text_block is None:
            raise ValueError(f"text 블록 없음 (content types: {[b.get('type') for b in resp.get('content', [])]})")
        text = text_block.strip()
        if text.startswith("```"):
            text = text.strip("`")
            if text.startswith("json"): text = text[4:]
        return json.loads(text), None
    except urllib.error.HTTPError as e:
        try:
            err_body = e.read().decode(errors="replace")[:200]
        except Exception:
            err_body = ""
        err = f"HTTP {e.code} {e.reason} - {err_body}"
        print(f"⚠️ LLM 판단 실패: {err}")
        return {}, err
    except Exception as e:
        err = f"{type(e).__name__}: {e}"
        print(f"⚠️ LLM 판단 실패: {err}")
        return {}, err

def build_report():
    rdate = datetime.now() - timedelta(days=1)
    dk = rdate.strftime("%Y-%m-%d")
    dlabel = rdate.strftime("%Y년 %m월 %d일")
    wd = ["월","화","수","목","금","토","일"][rdate.weekday()]
    dayoff = get_dayoff()
    now_str = datetime.now().strftime("%Y-%m-%d %H:%M")
    handover_recv = get_handover_confirmations(dk)  # 수신확인 성실도
    L = []

    L += [f"위베이프 일일 운영 보고",
          f"보고일: {dlabel} ({wd}요일)",
          f"생성: {now_str}", "","="*50,""]

    # 1. 매장별 현황
    L += ["[1] 매장별 운영 현황",""]
    issue_stores, ok_stores = [], []
    clean_pattern_flags = []

    for s in STORES:
        sid, sname, sshort = s["id"], s["name"], s["short"]
        cd, ct, ctimes = get_clean(sid, dk)
        clean_ok = ct>0 and cd>=ct
        pattern = check_clean_pattern(ctimes, ct)
        if pattern:
            clean_pattern_flags.append(f"{sshort}: {pattern}")
        hw = get_handover(sid, dk)
        unconf = [h for h in hw if not h.get("confirmed")]

        workers, td, tt = [], 0, 0
        for name in STAFF:
            if dayoff.get(name,{}).get(dk)=="dayoff": continue
            todos = [t for t in get_todos(name,dk) if t.get("storeId")==sid]
            if todos:
                workers.append(name); tt+=len(todos); td+=sum(1 for t in todos if t.get("done"))

        issues = []
        if not clean_ok and ct>0: issues.append(f"청소미완료({cd}/{ct})")
        if pattern: issues.append("청소체크 일괄처리 의심")
        if unconf: issues.append(f"인수인계미확인{len(unconf)}건")
        if not workers: issues.append("업무기록없음")

        if issues:
            L.append(f"  ⚠️  {sname}")
            for i in issues: L.append(f"      - {i}")
            if workers:
                cr = int(td/tt*100) if tt>0 else 0
                L.append(f"      업무 {td}/{tt}건({cr}%) | 근무: {', '.join(workers)}")
            if unconf:
                for u in unconf[:2]: L.append(f"      └ {u.get('text','')} (작성:{u.get('author','')})")
            issue_stores.append(sshort)
        else:
            cr = int(td/tt*100) if tt>0 else 100
            ok_stores.append(f"{sshort}({cr}%)")

    L.append("")
    if ok_stores: L.append(f"  ✅ 정상: {' / '.join(ok_stores)}")
    L += ["","="*50,""]

    # 2. 직원별 업무 현황
    L += ["[2] 직원별 업무 현황",""]
    active, dayoff_list, unused = [], [], []
    staff_texts = {}  # LLM 판단용 텍스트 수집

    for name in STAFF:
        if dayoff.get(name,{}).get(dk)=="dayoff":
            dayoff_list.append(name); continue
        todos = get_todos(name, dk)
        recv_cnt = handover_recv.get(name, 0)

        # 인수인계 텍스트는 업무 입력 여부와 무관하게 항상 수집
        my_handovers = []
        for s in STORES:
            my_handovers += [h.get("text","") for h in get_handover(s["id"], dk) if h.get("author")==name]

        if not todos:
            unused.append(name)
            if my_handovers:
                staff_texts[name] = {"todos": [], "handovers": my_handovers}
            continue

        total=len(todos); done=sum(1 for t in todos if t.get("done"))
        carried=sum(1 for t in todos if t.get("fromDate") and t.get("fromDate")!=dk)
        comp=int(done/total*100) if total>0 else 0
        bar="■"*(comp//10)+"□"*(10-comp//10)
        sids=list(set(t.get("storeId","") for t in todos if t.get("storeId")))
        snames=[next((s["short"] for s in STORES if s["id"]==sid),sid) for sid in sids]
        proj=[p for p in get_projects(name) if not p.get("done")]
        icon="✅" if comp>=80 else "⚠️" if comp>=50 else "❌"
        line=f"  {icon} {name:<4} [{bar}] {comp:3d}%  ({done}/{total}건)"
        if snames: line+=f"  @{','.join(snames)}"
        if carried: line+=f"  이월{carried}건"
        if proj: line+=f"  프로젝트{len(proj)}건"
        if recv_cnt: line+=f"  인계수신확인{recv_cnt}건"
        L.append(line)
        undone=[t for t in todos if not t.get("done")]
        for t in undone[:2]: L.append(f"          └❌ {t.get('text','')}")
        if len(undone)>2: L.append(f"          └ 외 {len(undone)-2}건")
        active.append({"name":name,"comp":comp})

        staff_texts[name] = {
            "todos": [t.get("text","") for t in todos],
            "handovers": my_handovers
        }

    L.append("")
    if dayoff_list: L.append(f"  🏖️ 휴무: {', '.join(dayoff_list)}")
    if unused:      L.append(f"  📵 미사용: {', '.join(unused)}")
    L += ["","="*50,""]

    # 3. AI 업무 내용 판단
    L += ["[3] AI 업무 내용 판단",""]
    has_content = any(d["todos"] or d["handovers"] for d in staff_texts.values())
    if not ANTHROPIC_API_KEY:
        L.append("  (ANTHROPIC_API_KEY 미설정 — 관리자 확인 필요)")
    elif not has_content:
        L.append("  (오늘은 업무·인수인계 입력이 없어 판단할 내용이 없습니다)")
    else:
        judgments, err = llm_judge(staff_texts)
        if judgments:
            for name in STAFF:
                if name in judgments:
                    L.append(f"  · {name}: {judgments[name]}")
        elif err:
            L.append(f"  (판단 실패: {err})")
        else:
            L.append("  (판단 실패 — 다음 실행에서 재시도됩니다)")
    L += ["","="*50,""]

    # 4. 인사이트
    L += ["[4] 오늘의 인사이트",""]
    insights = []
    if issue_stores: insights.append(f"• 이슈 매장 {len(issue_stores)}개: {', '.join(issue_stores)}")
    low=[s for s in active if s["comp"]<50]
    if low: insights.append(f"• 완료율 50% 미만: {', '.join(s['name'] for s in low)}")
    if unused: insights.append(f"• 앱 미사용 {len(unused)}명: {', '.join(unused)}")
    if clean_pattern_flags:
        insights.append(f"• 청소체크 일괄처리 의심: {' / '.join(clean_pattern_flags)}")
    if not insights: insights.append("• 특이사항 없음. 전반적으로 양호합니다.")
    L.extend(insights)
    L += ["","="*50,"📱 위베이프 운영 시스템 | 자동 발송"]
    return L, dlabel

def try_make_docx(lines, title):
    try:
        from docx import Document
        from docx.shared import Pt, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        doc = Document()
        for sec in doc.sections:
            sec.top_margin=Cm(2); sec.bottom_margin=Cm(2)
            sec.left_margin=Cm(2.5); sec.right_margin=Cm(2.5)
        h=doc.add_heading(title,0); h.alignment=WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph("")
        for line in lines:
            if line.startswith("="*10):
                p=doc.add_paragraph(); p.add_run("─"*40)
            elif line.startswith("[") and "]" in line[:5]:
                doc.add_heading(line, level=1)
            elif line == "":
                doc.add_paragraph("")
            else:
                p=doc.add_paragraph(line)
                p.paragraph_format.space_after=Pt(2)
        buf=io.BytesIO(); doc.save(buf); buf.seek(0)
        return buf.read()
    except ImportError:
        return None

def send():
    lines, dlabel = build_report()
    subject = f"[위베이프 일일보고] {dlabel}"
    body = "\n".join(lines)
    msg = MIMEMultipart()
    msg["From"]=GMAIL_USER; msg["To"]=REPORT_TO; msg["Subject"]=subject
    msg.attach(MIMEText(body, "plain", "utf-8"))
    docx = try_make_docx(lines, f"위베이프 일일보고 {dlabel}")
    if docx:
        part=MIMEBase("application","octet-stream"); part.set_payload(docx)
        encoders.encode_base64(part)
        fname=f"위베이프_일일보고_{dlabel.replace(' ','').replace('년','').replace('월','').replace('일','')}.docx"
        part.add_header("Content-Disposition", f"attachment; filename={fname}")
        msg.attach(part)
    with smtplib.SMTP("smtp.gmail.com", 587) as sv:
        sv.starttls(); sv.login(GMAIL_USER, GMAIL_APP_PASSWORD); sv.send_message(msg)
    print(f"✅ 발송 완료: {dlabel}")
    print(body)

if __name__=="__main__":
    send()
