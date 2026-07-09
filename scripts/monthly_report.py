"""
위베이프 월간 보고 스크립트 v4.0
- 기업 보고서 형식
- 6가지 역량 지표 (미완료 감점 없음)
- 베스트 직원/매장 TOP3
- Word(.docx) 첨부 발송
- 게시판 공지 자동 등록 (순위만, 점수 없음)
v4 추가:
  - 청소 체크 타임스탬프 기반 "몰아치기" 패턴 월간 집계
  - 인수인계 수신확인(다른 사람 인계를 내가 확인한 건수) → 협력도 점수에 반영
  - Claude API를 통한 업무/인수인계 내용 질적 판단 (직원별 월간 요약)
"""
import os, json, calendar, urllib.request, urllib.error, smtplib, io, time
from datetime import datetime, timedelta
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.base import MIMEBase
from email import encoders

FIREBASE_API_KEY   = os.environ["FIREBASE_API_KEY"]
GMAIL_USER         = os.environ["GMAIL_USER"]
GMAIL_APP_PASSWORD = os.environ["GMAIL_APP_PASSWORD"]
REPORT_TO          = os.environ["REPORT_TO_EMAIL"]
ANTHROPIC_API_KEY  = os.environ.get("ANTHROPIC_API_KEY")  # 없으면 LLM 판단 단계는 건너뜀
PROJECT_ID         = "wevape-schedule"
BASE_URL           = f"https://firestore.googleapis.com/v1/projects/{PROJECT_ID}/databases/(default)/documents"

STORES = [
    {"id":"yeonsu",    "name":"인천 연수점",   "short":"연수점"},
    {"id":"nonhyeon",  "name":"인천 논현점",   "short":"논현점"},
    {"id":"rodeo",     "name":"구월 로데오점", "short":"로데오점"},
    {"id":"gilbyeong", "name":"구월 길병원점", "short":"길병원점"},
    {"id":"airport",   "name":"인천공항점",    "short":"공항점"},
    {"id":"geomdan",   "name":"검단점",        "short":"검단점"},
    {"id":"gyesan",    "name":"계산점",        "short":"계산점"},
    {"id":"sangdong",  "name":"부천 상동점",   "short":"상동점"},
    {"id":"sijungdong","name":"부천 신중동점", "short":"신중동점"},
]

STAFF = ["오명록","고아현","장현진","장대운","신재현","정희경",
         "조효정","홍다운","이종혁","원주현","김형진","윤하람",
         "차영근","정유진","안태민","김다정"]

# ── Firebase 헬퍼 ───────────────────────────────
def fb_get(path):
    try:
        req = urllib.request.Request(f"{BASE_URL}/{path}?key={FIREBASE_API_KEY}")
        with urllib.request.urlopen(req, timeout=15) as r:
            d = json.load(r)
            return parse_fields(d["fields"]) if "fields" in d else None
    except: return None

def fb_patch(path, body):
    try:
        data = json.dumps(body).encode()
        req = urllib.request.Request(f"{BASE_URL}/{path}?key={FIREBASE_API_KEY}", data=data, method="PATCH")
        req.add_header("Content-Type","application/json")
        with urllib.request.urlopen(req, timeout=15) as r:
            return r.status == 200
    except: return False

def parse_value(v):
    if "stringValue"  in v: return v["stringValue"]
    if "booleanValue" in v: return v["booleanValue"]
    if "integerValue" in v: return int(v["integerValue"])
    if "doubleValue"  in v: return v["doubleValue"]
    if "arrayValue"   in v: return [parse_value(x) for x in v["arrayValue"].get("values",[])]
    if "mapValue"     in v: return parse_fields(v["mapValue"].get("fields",{}))
    return None

def parse_fields(f): return {k: parse_value(v) for k,v in f.items()}

def get_todos(name, dk):
    d = fb_get(f"staff_todos/{name}_{dk}"); return (d.get("items") or []) if d else []

# 청소 체크값 판독 (구버전 boolean / 신버전 {v,t} 객체 모두 지원)
def ck_val(x):
    if isinstance(x, dict): return x.get("v") is True
    return x is True
def ck_time(x):
    if isinstance(x, dict): return x.get("t")
    return None

def get_clean(sid, dk):
    d = fb_get(f"checks/{sid}/{dk}/clean")
    if not d: return 0, 0, []
    keys = [k for k in d if k.startswith("item_")]
    total = len(keys) or 6
    done = sum(1 for k in keys if ck_val(d.get(k)))
    times = sorted([ck_time(d.get(k)) for k in keys if ck_time(d.get(k))])
    return done, total, times

def check_clean_pattern(times, total):
    """청소 체크 타임스탬프가 비정상적으로 몰려서 찍혔는지(일괄 클릭 의심) 판단"""
    if len(times) < max(3, total // 2):
        return False
    span_sec = (max(times) - min(times)) / 1000
    return len(times) >= 4 and span_sec < 60

def get_handover(sid, dk):
    d = fb_get(f"handover/{sid}_{dk}"); return (d.get("items") or []) if d else []

def get_dayoff():
    d = fb_get("config/dayoff"); return d or {}
def get_projects(name):
    d = fb_get(f"staff_projects/{name}"); return (d.get("items") or []) if d else []

# ── 지난달 날짜 목록 ────────────────────────────
def get_prev_month_dates():
    today = datetime.now()
    first = today.replace(day=1)
    last_prev = first - timedelta(days=1)
    y, m = last_prev.year, last_prev.month
    num_days = calendar.monthrange(y, m)[1]
    dates = [f"{y}-{m:02d}-{d:02d}" for d in range(1, num_days+1)]
    return dates, y, m

# ── 인수인계 수신확인 집계 (월간, 전 매장) ───────
def build_handover_receipt_map(dates):
    """{이름: 그 달에 확인 처리한 인수인계 건수} - 수신 성실도"""
    counts = {}
    for dk in dates:
        for s in STORES:
            for it in get_handover(s["id"], dk):
                if it.get("confirmed") and it.get("confirmedBy"):
                    who = it["confirmedBy"]
                    counts[who] = counts.get(who, 0) + 1
    return counts

# ── 직원 분석 ───────────────────────────────────
def analyze_staff(name, dates, dayoff_data, handover_recv_map):
    my_dayoffs = dayoff_data.get(name, {})
    work_days, used_days = 0, 0
    total_todos, done_todos, kept_count, carry_count, carry_done = 0,0,0,0,0
    handover_written, handover_confirmed = 0,0
    sample_todos, sample_handovers = [], []

    for dk in dates:
        is_dayoff = my_dayoffs.get(dk) == "dayoff"
        if is_dayoff: continue
        work_days += 1

        todos = get_todos(name, dk)
        if todos:
            used_days += 1
            total_todos += len(todos)
            done_todos += sum(1 for t in todos if t.get("done"))
            for t in todos:
                if t.get("fromDate") and t.get("fromDate") != dk:
                    carry_count += 1
                    if t.get("done"): carry_done += 1
                    if t.get("kept"): kept_count += 1
            if len(sample_todos) < 20:
                sample_todos += [t.get("text","") for t in todos][:3]

    # 인수인계 확인률 (내가 쓴 것이 확인됐는지 = 발신측)
    for s in STORES:
        for dk in dates:
            items = get_handover(s["id"], dk)
            mine = [i for i in items if i.get("author")==name]
            handover_written += len(mine)
            handover_confirmed += sum(1 for i in mine if i.get("confirmed"))
            if len(sample_handovers) < 10:
                sample_handovers += [i.get("text","") for i in mine][:2]

    # 인수인계 수신확인 건수 (내가 남의 것을 확인한 = 수신측, 성실도)
    handover_received = handover_recv_map.get(name, 0)

    # 프로젝트
    projs = get_projects(name)
    ongoing = len([p for p in projs if not p.get("done")])
    completed_p = len([p for p in projs if p.get("done")])

    # 6가지 역량 점수
    part_rate = used_days/work_days if work_days>0 else 0
    comp_rate = done_todos/total_todos if total_todos>0 else 0
    carry_resolved = carry_done + kept_count
    carry_rate = carry_resolved/carry_count if carry_count>0 else 1.0
    handover_send_rate = handover_confirmed/handover_written if handover_written>0 else 1.0
    # 수신 성실도 보정치: 근무일 대비 수신확인 건수 (하루 1건 이상이면 만점 취급)
    handover_recv_rate = min(handover_received/work_days, 1.0) if work_days>0 else 0

    s_steady     = int(part_rate * 20)                        # 꾸준함 20점
    s_diligence  = min(total_todos//3, 12) + int(comp_rate*13)# 충실도 25점
    s_resp       = int(carry_rate * 20)                       # 책임감 20점
    s_initiative = min(ongoing*4 + completed_p*2, 10)         # 주도성 10점
    # 협력도 15점 = 발신측(내가 남긴 인계가 확인됐는지) 10점 + 수신측(내가 확인 처리했는지) 5점
    s_handover   = int(handover_send_rate*10) + int(handover_recv_rate*5)
    s_growth     = 10                                         # 성장성 10점 (기본값)
    total_score  = s_steady+s_diligence+s_resp+s_initiative+s_handover+s_growth

    # 유형 분류
    if used_days == 0:                                stype = "📵 미참여형"
    elif s_initiative>=8 and part_rate>=0.8:          stype = "🌟 주도형"
    elif part_rate>=0.85 and carry_rate>=0.8:         stype = "✅ 성실형"
    elif total_todos>=10 and comp_rate<0.5 and carry_rate<0.5: stype = "📋 형식형"
    elif part_rate<=0.3:                              stype = "📵 미참여형"
    else:                                             stype = "📌 일반형"

    return {
        "name":name, "score":total_score, "type":stype,
        "used_days":used_days, "work_days":work_days,
        "total":total_todos, "done":done_todos, "comp_rate":comp_rate,
        "carry_count":carry_count, "carry_done":carry_done,
        "carry_rate":carry_rate, "kept_count":kept_count,
        "ongoing":ongoing, "completed_p":completed_p,
        "handover_rate":handover_send_rate,
        "handover_received":handover_received,
        "handover_recv_rate":handover_recv_rate,
        "sample_todos":sample_todos, "sample_handovers":sample_handovers,
        "scores":{"꾸준함":s_steady,"충실도":s_diligence,"책임감":s_resp,
                  "주도성":s_initiative,"협력도":s_handover,"성장성":s_growth}
    }

# ── 매장 분석 ───────────────────────────────────
def analyze_store(store, dates, dayoff_data):
    sid = store["id"]
    clean_days=clean_total_days=0
    todo_total=todo_done=0
    hw_total=hw_confirmed=0
    pattern_days=0

    for dk in dates:
        cd,ct,ctimes = get_clean(sid, dk)
        if ct>0:
            clean_total_days+=1
            if cd>=ct: clean_days+=1
            if check_clean_pattern(ctimes, ct): pattern_days+=1

        for name in STAFF:
            if dayoff_data.get(name,{}).get(dk)=="dayoff": continue
            todos=[t for t in get_todos(name,dk) if t.get("storeId")==sid]
            todo_total+=len(todos)
            todo_done+=sum(1 for t in todos if t.get("done"))

        items=get_handover(sid,dk)
        hw_total+=len(items)
        hw_confirmed+=sum(1 for i in items if i.get("confirmed"))

    clean_rate = clean_days/clean_total_days if clean_total_days>0 else 0
    todo_rate  = todo_done/todo_total if todo_total>0 else 0
    hw_rate    = hw_confirmed/hw_total if hw_total>0 else 1.0
    score      = int(clean_rate*40 + todo_rate*35 + hw_rate*25)

    return {
        "name":store["name"],"short":store["short"],"score":score,
        "clean_rate":clean_rate,"clean_days":clean_days,"clean_total":clean_total_days,
        "todo_rate":todo_rate,"todo_done":todo_done,"todo_total":todo_total,
        "hw_rate":hw_rate,"pattern_days":pattern_days
    }

# ── Claude API를 통한 업무/인수인계 내용 질적 판단 (월간 요약) ──
def llm_judge_monthly(staff_results):
    """활동한 직원들만 대상으로, 한 달 샘플 텍스트 기반 한 줄 질적 평가. 반환: (dict, error or None)"""
    if not ANTHROPIC_API_KEY:
        return {}, None
    lines = []
    for s in staff_results:
        if s["used_days"] == 0 and not s["sample_handovers"]:
            continue
        name = s["name"]
        lines.append(f"[{name}] (근무 {s['used_days']}일, 업무 {s['total']}건)")
        if s["sample_todos"]:
            lines.append("업무 샘플: " + " / ".join(s["sample_todos"][:12]))
        if s["sample_handovers"]:
            lines.append("인수인계 샘플: " + " / ".join(s["sample_handovers"][:8]))
    if not lines:
        return {}, None

    prompt = (
        "다음은 위베이프 매장 직원들이 한 달 동안 입력한 '업무 내용'과 '인수인계 내용'의 샘플입니다.\n"
        "각 직원별로 입력 내용이 한 달 동안 대체로 구체적이고 실질적이었는지, 형식적이거나 성의없이 "
        "짧게만 썼는지를 한두 문장으로 담백하게 평가해 주세요. 과장하지 말고 샘플에 근거해서만 판단하세요.\n"
        "반드시 아래 JSON 형식으로만 응답하세요 (다른 설명, 코드블록 없이 순수 JSON만):\n"
        '{"직원이름": "한두 문장 평가"}\n\n'
        + "\n".join(lines)
    )
    try:
        body = json.dumps({
            "model": "claude-sonnet-5",
            "max_tokens": 2000,
            "messages": [{"role": "user", "content": prompt}]
        }).encode()
        req = urllib.request.Request(
            "https://api.anthropic.com/v1/messages",
            data=body, method="POST"
        )
        req.add_header("x-api-key", ANTHROPIC_API_KEY)
        req.add_header("anthropic-version", "2023-06-01")
        req.add_header("content-type", "application/json")
        with urllib.request.urlopen(req, timeout=45) as r:
            resp = json.load(r)
        text_block = next((b.get("text") for b in resp.get("content", []) if b.get("type") == "text"), None)
        if text_block is None:
            raise ValueError(f"text 블록 없음 (content types: {[b.get('type') for b in resp.get('content', [])]})")
        text = text_block.strip()
        if text.startswith("```"):
            text = text.strip("`")
            if text.startswith("json"): text = text[4:]
        return json.loads(text), None
    except urllib.error.HTTPError as e:
        try:
            err_body = e.read().decode(errors="replace")[:200]
        except Exception:
            err_body = ""
        err = f"HTTP {e.code} {e.reason} - {err_body}"
        print(f"⚠️ LLM 월간 판단 실패: {err}")
        return {}, err
    except Exception as e:
        err = f"{type(e).__name__}: {e}"
        print(f"⚠️ LLM 월간 판단 실패: {err}")
        return {}, err

# ── 게시판 공지 등록 ────────────────────────────
def post_board_notice(y, m, best_staff, best_stores):
    month_label = f"{y}년 {m}월"
    medals = ["🥇","🥈","🥉"]
    body = f"{month_label} 한 달 수고 많으셨습니다!\n\n"
    body += "🏆 이달의 베스트 직원\n"
    for i,s in enumerate(best_staff[:3]):
        body += f"{medals[i]} {s['name']}\n"
    body += "\n🏆 이달의 베스트 매장\n"
    for i,s in enumerate(best_stores[:3]):
        body += f"{medals[i]} {s['name']}\n"
    body += "\n모두 수고 많으셨습니다 👏"

    existing = fb_get("board/posts")
    items = (existing.get("items") or []) if existing else []
    now = datetime.now()
    new_post = {
        "id": str(int(time.time()*1000)), "cat":"notice",
        "title": f"🏆 {month_label} 이달의 베스트",
        "body": body, "author":"관리자",
        "date": f"{now.month}/{now.day} {now.hour:02d}:{now.minute:02d}"
    }
    def to_v(v):
        if isinstance(v,bool): return {"booleanValue":v}
        return {"stringValue":str(v)}
    def to_map(d): return {"mapValue":{"fields":{k:to_v(v) for k,v in d.items()}}}
    fb_body={"fields":{"items":{"arrayValue":{"values":[to_map(p) for p in [new_post]+items]}}}}
    return fb_patch("board/posts", fb_body)

# ── 보고서 생성 ─────────────────────────────────
def build_report():
    dates, y, m = get_prev_month_dates()
    month_label = f"{y}년 {m}월"
    dayoff = get_dayoff()
    now_str = datetime.now().strftime("%Y-%m-%d %H:%M")
    medals = ["🥇","🥈","🥉"]
    L = []

    L += [f"위베이프 월간 운영 보고",
          f"보고 기간: {month_label}",
          f"생성: {now_str}","","="*55,""]

    # 1. 경영진 요약
    L += ["[1] 이달의 경영 요약",""]

    print("인수인계 수신확인 집계 중...")
    handover_recv_map = build_handover_receipt_map(dates)

    print("직원 데이터 수집 중...")
    staff_results = []
    for name in STAFF:
        print(f"  - {name}...")
        r = analyze_staff(name, dates, dayoff, handover_recv_map)
        staff_results.append(r)

    print("매장 데이터 수집 중...")
    store_results = []
    for s in STORES:
        print(f"  - {s['name']}...")
        r = analyze_store(s, dates, dayoff)
        store_results.append(r)

    staff_sorted = sorted(staff_results, key=lambda x:x["score"], reverse=True)
    store_sorted = sorted(store_results, key=lambda x:x["score"], reverse=True)
    active = [s for s in staff_sorted if s["used_days"]>0]

    total_todos_all = sum(s["total"] for s in staff_results)
    total_done_all  = sum(s["done"] for s in staff_results)
    overall_comp    = int(total_done_all/total_todos_all*100) if total_todos_all>0 else 0
    no_show = [s for s in staff_results if s["used_days"]==0]
    leaders = [s for s in staff_results if s["type"]=="🌟 주도형"]
    pattern_stores = [s for s in store_results if s["pattern_days"]>0]

    L.append(f"  전체 업무 완료율: {overall_comp}% ({total_done_all}/{total_todos_all}건)")
    L.append(f"  앱 참여 직원: {len(active)}/{len(STAFF)}명")
    if leaders: L.append(f"  이달의 주도형 직원: {', '.join(s['name'] for s in leaders)}")
    if no_show: L.append(f"  ⚠️ 앱 미참여: {', '.join(s['name'] for s in no_show)}")
    if pattern_stores:
        pat_str = ', '.join(f"{s['short']}({s['pattern_days']}일)" for s in pattern_stores)
        L.append(f"  ⚠️ 청소체크 일괄처리 의심 매장: {pat_str}")
    L += ["","="*55,""]

    # 2. 매장별 월간 현황
    L += ["[2] 매장별 월간 운영 현황",""]
    for s in store_sorted:
        bar = "■"*(s["score"]//10) + "□"*(10-s["score"]//10)
        L.append(f"  {s['name']}  [{bar}] {s['score']}점")
        L.append(f"    청소완료율  {int(s['clean_rate']*100):3d}%  ({s['clean_days']}/{s['clean_total']}일)")
        if s["pattern_days"]>0:
            L.append(f"    ⚠️ 청소체크 일괄처리 의심 {s['pattern_days']}일")
        L.append(f"    업무완료율  {int(s['todo_rate']*100):3d}%  ({s['todo_done']}/{s['todo_total']}건)")
        L.append(f"    인수인계확인 {int(s['hw_rate']*100):3d}%")
        L.append("")
    L += ["="*55,""]

    # 3. 직원별 역량 분석 (관리자 전용)
    L += ["[3] 직원별 역량 분석  ※ 관리자 전용","",
          "  ※ 미완료 업무 자체는 감점 없음 — 이월 후 처리 여부로 책임감 평가",
          "  ※ 협력도(15점) = 인계 발신 확인률(10점) + 인계 수신확인 성실도(5점)",""]

    for s in staff_sorted:
        bar = "■"*(s["score"]//10) + "□"*(10-s["score"]//10)
        comp = int(s["comp_rate"]*100)
        part = int(s["used_days"]/s["work_days"]*100) if s["work_days"]>0 else 0
        L.append(f"  {s['type']}  {s['name']}  [{bar}]  {s['score']}점")
        L.append(f"    참여율  {part}%  ({s['used_days']}/{s['work_days']}일 근무)")
        L.append(f"    업무    {s['total']}건 입력 → {s['done']}건 완료 ({comp}%)")
        if s["carry_count"]>0:
            txt = f"    이월    {s['carry_count']}건 중 처리 {s['carry_done']}건"
            if s["kept_count"]>0: txt += f" / 보관 {s['kept_count']}건(합리적사유)"
            L.append(txt)
        L.append(f"    인수인계  발신확인률 {int(s['handover_rate']*100)}% · 수신확인 {s['handover_received']}건")
        if s["ongoing"] or s["completed_p"]:
            L.append(f"    프로젝트 진행중 {s['ongoing']}건 · 완료 {s['completed_p']}건")
        sc = s["scores"]
        L.append(f"    세부: 꾸준함{sc['꾸준함']} 충실도{sc['충실도']} 책임감{sc['책임감']} 주도성{sc['주도성']} 협력도{sc['협력도']} 성장성{sc['성장성']}")
        L.append("")
    L += ["="*55,""]

    # 4. AI 업무 내용 판단 (월간)
    L += ["[4] AI 업무 내용 판단 (월간 샘플 기반)",""]
    has_content = any(s["sample_todos"] or s["sample_handovers"] for s in staff_results)
    if not ANTHROPIC_API_KEY:
        L.append("  (ANTHROPIC_API_KEY 미설정 — 관리자 확인 필요)")
    elif not has_content:
        L.append("  (이번 달 업무·인수인계 입력이 없어 판단할 내용이 없습니다)")
    else:
        judgments, err = llm_judge_monthly(staff_results)
        if judgments:
            for s in staff_sorted:
                if s["name"] in judgments:
                    L.append(f"  · {s['name']}: {judgments[s['name']]}")
        elif err:
            L.append(f"  (판단 실패: {err})")
        else:
            L.append("  (판단 실패 — 다음 실행에서 재시도됩니다)")
    L += ["","="*55,""]

    # 5. 이달의 베스트 TOP3
    L += ["[5] 이달의 베스트",""]
    L.append(f"  🏆 베스트 직원 TOP3")
    for i,s in enumerate(active[:3]):
        comp=int(s["comp_rate"]*100)
        L.append(f"    {medals[i]} {s['name']}  {s['type']}  {s['score']}점")
        L.append(f"         참여율 {int(s['used_days']/s['work_days']*100) if s['work_days']>0 else 0}% · 업무완료 {comp}% · 프로젝트 {s['ongoing']+s['completed_p']}건")
    L.append("")
    L.append(f"  🏆 베스트 매장 TOP3")
    for i,s in enumerate(store_sorted[:3]):
        L.append(f"    {medals[i]} {s['name']}  {s['score']}점")
        L.append(f"         청소 {int(s['clean_rate']*100)}% · 업무 {int(s['todo_rate']*100)}% · 인수인계 {int(s['hw_rate']*100)}%")
    L += ["","="*55,""]

    # 6. 종합 의견
    L += ["[6] 월간 종합 의견",""]
    opinions = []
    low_stores = [s for s in store_results if s["clean_rate"]<0.8]
    if low_stores: opinions.append(f"• 청소 수행률 부진 매장: {', '.join(s['short'] for s in low_stores)} — 점검 필요")
    if pattern_stores: opinions.append(f"• 청소체크 일괄처리 의심 매장: {', '.join(s['short'] for s in pattern_stores)} — 현장 확인 권장")
    low_carry = [s for s in staff_results if s["carry_count"]>=5 and s["carry_rate"]<0.4]
    if low_carry: opinions.append(f"• 이월 미처리 누적: {', '.join(s['name'] for s in low_carry)} — 업무 부하 점검")
    low_recv = [s for s in staff_results if s["used_days"]>=5 and s["handover_recv_rate"]<0.2]
    if low_recv: opinions.append(f"• 인수인계 수신확인 저조: {', '.join(s['name'] for s in low_recv)} — 인계 확인 습관 점검 필요")
    if no_show: opinions.append(f"• 앱 미참여 직원 {len(no_show)}명 — 현장 확인 필요")
    if leaders: opinions.append(f"• 우수 활약: {', '.join(s['name'] for s in leaders)} — 사례 공유 권장")
    if not opinions: opinions.append("• 전반적으로 양호한 한 달이었습니다.")
    L.extend(opinions)
    L += ["","="*55,"📱 위베이프 운영 시스템 | 자동 월간 보고"]

    return L, month_label, y, m, active, store_sorted

def try_make_docx(lines, title):
    try:
        from docx import Document
        from docx.shared import Pt, Cm, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        doc = Document()
        for sec in doc.sections:
            sec.top_margin=Cm(2); sec.bottom_margin=Cm(2)
            sec.left_margin=Cm(2.5); sec.right_margin=Cm(2.5)
        h = doc.add_heading(title, 0); h.alignment=WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph("")
        for line in lines:
            if line.startswith("="*10):
                p=doc.add_paragraph(); p.add_run("─"*45)
            elif line.startswith("[") and "]" in line[:5]:
                doc.add_heading(line, level=1)
            elif line=="":
                doc.add_paragraph("")
            else:
                p=doc.add_paragraph(line)
                p.paragraph_format.space_after=Pt(1)
        buf=io.BytesIO(); doc.save(buf); buf.seek(0)
        return buf.read()
    except ImportError:
        return None

def send():
    lines, month_label, y, m, active, store_sorted = build_report()
    subject = f"[위베이프 월간보고] {month_label}"
    body = "\n".join(lines)
    msg = MIMEMultipart()
    msg["From"]=GMAIL_USER; msg["To"]=REPORT_TO; msg["Subject"]=subject
    msg.attach(MIMEText(body,"plain","utf-8"))
    docx = try_make_docx(lines, f"위베이프 월간보고 {month_label}")
    if docx:
        part=MIMEBase("application","octet-stream"); part.set_payload(docx)
        encoders.encode_base64(part)
        fname=f"위베이프_월간보고_{month_label.replace(' ','')}.docx"
        part.add_header("Content-Disposition",f"attachment; filename={fname}")
        msg.attach(part)
    with smtplib.SMTP("smtp.gmail.com",587) as sv:
        sv.starttls(); sv.login(GMAIL_USER,GMAIL_APP_PASSWORD); sv.send_message(msg)
    posted = post_board_notice(y, m, active, store_sorted)
    print(f"✅ 월간보고 발송 완료: {month_label}")
    print(f"{'✅' if posted else '⚠️'} 게시판 공지 {'등록 완료' if posted else '등록 실패'}")
    print(body)

if __name__=="__main__":
    send()
